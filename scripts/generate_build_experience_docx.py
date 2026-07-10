#!/usr/bin/env python3
"""Generate Viral_Velocity_App_Build_Experience.docx — narrative experience document."""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "Viral_Velocity_App_Build_Experience.docx"


def heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def para(doc, text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    return p


def bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_table(doc, headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
        for r in t.rows[0].cells[i].paragraphs[0].runs:
            r.bold = True
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            t.rows[ri + 1].cells[ci].text = str(val)
    doc.add_paragraph()


def build():
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    title = doc.add_heading("Building Viral Velocity", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph("Our Experience Building an AI-Assisted Mobile App")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].bold = True
    sub.runs[0].font.size = Pt(13)
    meta = doc.add_paragraph("Project: Viral Velocity Engine  |  June 2026")
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    heading(doc, "Introduction")
    para(doc,
        "This document shares what it was like to build the Viral Velocity app using AI coding "
        "tools (Cursor). It focuses on experience — what worked, what broke, and what we learned — "
        "not technical specifications. We started from a Business Requirements Document and a 5-day "
        "plan. We delivered a working backend, admin portal, and mobile app tested on a real Android phone."
    )

    heading(doc, "How We Worked")
    para(doc,
        "We treated the AI as a fast junior developer: describe the feature, generate code, run it, "
        "report bugs in plain language, fix, repeat. A human was still needed for product decisions, "
        "security requirements, physical device testing, and budget choices (free vs paid APIs)."
    )

    heading(doc, "The Journey")
    heading(doc, "Days 1–2: Backend & Admin — “This is going fast”", 2)
    para(doc,
        "Express, SQLite, auth, photo upload, scoring, enhancement, subscriptions, and the admin "
        "portal came together quickly. Database migrations matched the BRD. Admin dashboard, subscribers, "
        "discounts, and 2FA were usable with minimal hand coding."
    )
    para(doc, "We had to explicitly ask for:", bold=True)
    bullets(doc, [
        "OTP rate limiting and brute-force protection",
        "Forgot-password responses that do not reveal if an email exists",
        "Proper errors when uploading more than 5 photos",
    ])
    para(doc, "AI effectiveness this phase: 4–5 / 5", italic=True)

    heading(doc, "Days 3–4: Mobile — “Pretty screens, harder interactions”", 2)
    para(doc,
        "Seventeen screens with NativeWind styling. Auth, upload, scoring, library, and subscription "
        "layouts were generated in good structure. The hardest feature was swipe-to-save/discard on "
        "enhanced photos — gesture-handler and Reanimated needed several tuning rounds. Android bundling "
        "failed once until babel-preset-expo was installed."
    )
    para(doc, "AI effectiveness this phase: 3–4 / 5", italic=True)

    heading(doc, "Day 5: QA — “Proof on paper”", 2)
    para(doc,
        "30/30 automated backend checks passed. Part 2 logs were filled. We thought we were done — "
        "but passing API tests is not the same as the app feeling right on a phone."
    )

    heading(doc, "After delivery: Real phone testing — “Where reality hit”", 2)
    para(doc,
        "Testing on Android via Expo Go over Wi-Fi revealed a second wave of issues: status bar overlap, "
        "white system navigation bar, LAN API URL, prompts failing (Replicate needed paid credits), "
        "subscription UI not matching wireframe, app crash from outdated Expo APIs, and hidden settings."
    )
    para(doc,
        "We switched enhancement to free local Sharp processing with a keyword parser. Prompts like "
        "“warm golden sunset” work; “add a hat” cannot — and we now tell users honestly."
    )
    para(doc, "AI effectiveness this phase: 3 / 5 alone, 4 / 5 with clear bug reports from real use", italic=True)

    heading(doc, "Where AI Helped Most")
    bullets(doc, [
        "Boilerplate, CRUD, migrations, JWT auth",
        "Admin web forms, tables, React Query",
        "Consistent patterns across the codebase",
        "Documentation and QA scripts",
        "Fast fixes when given exact error messages",
    ])
    para(doc, "Estimate: AI handled roughly 70–80% of structure and typing.", italic=True)

    heading(doc, "Where AI Struggled")
    bullets(doc, [
        "Security features not added unless explicitly requested",
        "Mobile gestures and swipe interactions",
        "Device-specific UI (safe areas, system bars, LAN IP)",
        "Paid API economics (Replicate vs free Sharp)",
        "BRD contradictions (token tiers vs Monthly/Annual plans)",
        "Outdated API knowledge (Expo SDK 56 breaking changes)",
    ])

    heading(doc, "Lessons Learned")
    bullets(doc, [
        "AI is a pair programmer, not autopilot.",
        "Test on a physical device before calling a feature done.",
        "Ask explicitly for security and edge cases.",
        "When the BRD conflicts with itself, a human must choose.",
        "Free local tools can replace paid AI if expectations are set clearly.",
    ])

    heading(doc, "Did We Succeed?")
    para(doc,
        "Yes. We delivered a full monorepo: backend (30/30 QA), admin web, and mobile app on Expo Go. "
        "The last mile — security, gestures, device polish, and enhancement realism — required human "
        "testing and product decisions."
    )

    heading(doc, "Self-Ratings (AI Effectiveness, 1–5)")
    add_table(doc,
        ["Area", "Score", "Reason"],
        [
            ("Backend APIs", "5", "Fast, accurate, easy to extend"),
            ("Admin web", "5", "Standard React; few surprises"),
            ("Mobile UI (static)", "4", "Good layouts; minor polish"),
            ("Mobile gestures", "3", "Several fix rounds"),
            ("Real-device QA", "3", "Issues found only on phone"),
            ("Enhancement / prompts", "3", "Pivot to Sharp after API costs"),
            ("Documentation", "5", "Easy to generate and refine"),
            ("Overall", "4", "Strong accelerator; needs human judgment"),
        ],
    )

    heading(doc, "One-Sentence Summary")
    p = doc.add_paragraph()
    p.add_run(
        "AI got us to a demo-ready app quickly; human testing, product decisions, and honest limits "
        "on “AI enhancement” turned it into something we could show and explain with confidence."
    ).italic = True

    heading(doc, "Note for Management")
    para(doc,
        "This assignment evaluated AI coding tools on a production-style app. Our conclusion: they build "
        "most of the application well, but the final mile — security, mobile UX on real hardware, paid "
        "integrations, and ambiguous requirements — still requires an experienced developer in the loop."
    )

    doc.add_paragraph()
    end = doc.add_paragraph("— End of experience document —")
    end.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(OUT)
    print(f"Created: {OUT}")


if __name__ == "__main__":
    build()
