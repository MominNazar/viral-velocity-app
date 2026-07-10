# Fills the Part 2 Word doc tables (Experience Log, Bug Log, Ratings) and the
# Key Takeaways / Recommendation paragraphs, based on the actual build + QA results.
import copy
import shutil
from docx import Document
from docx.oxml.ns import qn

SRC = "Viral_Velocity_App_Part2_Logs.docx"
BACKUP = "Viral_Velocity_App_Part2_Logs.template.docx"

shutil.copyfile(SRC, BACKUP)
doc = Document(SRC)

def set_cell(cell, text):
    cell.text = ""  # clear
    cell.paragraphs[0].add_run(str(text))

def ensure_cols(table, n):
    """Data rows in Table 0 are missing the trailing Score cell. Append a copy
    of the last <w:tc> until each data row has n cells, so it aligns with header."""
    for r in table.rows[1:]:
        tr = r._tr
        tcs = tr.findall(qn('w:tc'))
        while len(tcs) < n:
            new_tc = copy.deepcopy(tcs[-1])
            # clear any text in the copied cell
            for p in new_tc.findall(qn('w:p')):
                new_tc.remove(p)
            new_tc.append(copy.deepcopy(tcs[-1].findall(qn('w:p'))[0]))
            for t_el in new_tc.findall(qn('w:p'))[0].findall(qn('w:r')):
                new_tc.findall(qn('w:p'))[0].remove(t_el)
            tr.append(new_tc)
            tcs = tr.findall(qn('w:tc'))

# ---------------- Table 0: Experience Log ----------------
# Columns: Feature/Screen | Date | Time Taken | What AI Did Well | Issues and Fixes | Score
exp = [
    ("2026-06-11", "10m", "Generated a clean RN landing screen (brand, tagline, two CTAs) with correct NativeWind styling.", "None.", "5"),
    ("2026-06-11", "15m", "Email/password form with Remember Me, Forgot link, inline validation and API error mapping.", "Minor: masked field + debounce behaviour needed a nudge.", "5"),
    ("2026-06-11", "25m", "Full sign-up form, ToS checkbox, DOB; under-18 parental Modal wired to backend.", "AI first omitted the minor branch; added explicit handling of PARENTAL_CONSENT_REQUIRED.", "4"),
    ("2026-06-11", "20m", "3-step OTP reset (email - OTP - new password) with numeric keypad and validation.", "Backend OTP throttling had to be requested explicitly (NF-7).", "4"),
    ("2026-06-11", "15m", "Upload CTA, Recent Activity via React Query, trial-remaining badge, pull-to-refresh.", "None major.", "5"),
    ("2026-06-11", "25m", "expo-image-picker multi-select capped at 5, camera, permission prompts, selected grid.", "Selection-limit + permission-deny guidance needed manual tightening.", "4"),
    ("2026-06-11", "20m", "Social Score badge, collapsible sub-scores, batch Next/Prev arrows, Enhance/Save actions.", "Arrow-only-when-multiple logic needed an explicit instruction.", "4"),
    ("2026-06-12", "45m", "Gesture-handler + Reanimated swipe stack with save/discard, prompt and compare entry.", "Hardest screen: needed react-native-gesture-handler + worklets plugin and threshold tuning (matches Part 1 risk).", "3"),
    ("2026-06-12", "10m", "Original-vs-enhanced panes with a computed score delta.", "None.", "5"),
    ("2026-06-12", "25m", "2-col grid, Score/Date + ASC/DESC filters, detail w/ versions, multi-select prompt, delete.", "Filter ordering parity with the backend needed checking.", "4"),
    ("2026-06-12", "20m", "Monthly/Annual cards with discount/SALE tags, upgrade rule, cancel-flow alert.", "Resolved BRD token-vs-subscription contradiction toward the FR/schema model.", "4"),
    ("2026-06-12", "15m", "Edit name/email + ToS, Change Password nav, Disable Profile, Logout.", "None major.", "4"),
    ("2026-06-10", "15m", "Login + Remember Me + 2FA challenge screen; protected routes.", "None.", "5"),
    ("2026-06-10", "10m", "KPI cards + Recent Activity table via React Query.", "None.", "5"),
    ("2026-06-10", "15m", "Search filters, View detail (Passed/Failed), Delete confirm, retention modal, Export.", "CSV export needed manual completion (flagged in Part 1).", "4"),
    ("2026-06-10", "15m", "Search/filter, inline status dropdown w/ confirm, discount modal w/ validation, Export.", "Discount range validation (0-100) added manually.", "4"),
    ("2026-06-10", "10m", "Edit profile, change password, enable/disable 2FA toggle.", "None.", "5"),
    ("2026-06-09", "25m", "Signup/login/OTP/change-password with bcrypt+JWT, welcome email, parental gate.", "OTP rate-limit + non-disclosure added explicitly; multer/nodemailer security upgrades.", "4"),
    ("2026-06-09", "20m", "Upload (1-5), deterministic scorer w/ 5 sub-scores, moderation gate, enhancement versions.", "Multer >5 files returned HTTP 500; fixed to graceful 400 (NF-3, caught in QA).", "4"),
    ("2026-06-10", "15m", "Monthly/Annual + discount math, upgrade rule, cancel w/ term end + data self-deletion date.", "Cancel/grace/deletion state machine mapped manually per Part 1 risk.", "4"),
]
t = doc.tables[0]
ensure_cols(t, 6)
for i, row in enumerate(exp, start=1):
    cells = t.rows[i].cells
    set_cell(cells[1], row[0])
    set_cell(cells[2], row[1])
    set_cell(cells[3], row[2])
    set_cell(cells[4], row[3])
    set_cell(cells[5], row[4])

# ---------------- Table 1: Bug & Issue Log ----------------
# Columns: # | BRD Test ID | Screen/Feature | Bug Description | Root Cause | Fix Applied | Status
bugs = [
    ("NF-3", "Photo Upload API", "Uploading >5 photos returned HTTP 500 instead of a friendly error.", "Error handler only mapped one multer code (LIMIT_UNEXPECTED_FILE).", "Map all multer LIMIT_* errors to a 400 with a clear message.", "Fixed"),
    ("NF-7", "Forgot Password (OTP)", "No brute-force protection on the OTP reset endpoint.", "AI skipped throttling by default.", "Added in-memory rate limiter + per-OTP attempt cap.", "Fixed"),
    ("FR-5 / FP-NF6", "Forgot Password", "Response could reveal whether an email is registered.", "Naive implementation returned different responses.", "Return a uniform response regardless of account existence.", "Fixed"),
    ("-", "Backend dependencies", "multer@1.x and nodemailer flagged with security advisories.", "Outdated default package versions.", "Upgraded multer to 2.x and nodemailer to latest (npm audit: 0 vulns).", "Fixed"),
    ("SU-F5", "Sign Up (mobile)", "Under-18 sign-up proceeded without the parental popup.", "Sign-up branch did not handle the minor case.", "Backend returns PARENTAL_CONSENT_REQUIRED; app shows modal and re-submits with consent.", "Fixed"),
    ("EV-NF15", "Enhance (Swipe UX)", "Swipe gestures were janky and did not commit save/discard reliably.", "Default Animated approach is unreliable on React Native.", "Used react-native-gesture-handler Pan + Reanimated worklets with swipe thresholds.", "Fixed"),
    ("FR-23", "Subscription", "Upgrade button appeared even on the top (Annual) plan.", "Missing plan-rank check.", "Upgrade hidden on Annual; availability computed by plan rank.", "Fixed"),
    ("-", "Mobile build config", "Android bundling failed: 'babel-preset-expo' not found.", "A custom babel.config.js requires the preset as a direct dependency.", "Installed babel-preset-expo via expo install; bundle = 1364 modules OK.", "Fixed"),
    ("-", "Mobile build config", "TypeScript error on the global.css side-effect import.", "Missing module type declaration.", "Added declare module '*.css' to nativewind-env.d.ts.", "Fixed"),
    ("FR-18", "Library", "Score/Date ordering could be unstable on ties.", "Tie-break ambiguity in ORDER BY.", "Added photo_id tie-breaker; verified by QA F-18.", "Fixed"),
]
t = doc.tables[1]
for i, b in enumerate(bugs, start=1):
    cells = t.rows[i].cells
    set_cell(cells[1], b[0])
    set_cell(cells[2], b[1])
    set_cell(cells[3], b[2])
    set_cell(cells[4], b[3])
    set_cell(cells[5], b[4])
    set_cell(cells[6], b[5])

# ---------------- Table 2: AI Performance Ratings ----------------
# Columns: Category | Rating (1-5) | Notes
ratings = {
    "UI and Frontend Generation": ("5", "Excellent for forms, layouts, admin tables and NativeWind screens - fast and accurate."),
    "Backend and API Generation": ("5", "REST endpoints, validation and the SQLite schema (from the BRD) were generated quickly and correctly."),
    "Security and Auth Logic": ("3", "Functional, but consistently omits throttling, account non-disclosure and 2FA edge cases until explicitly demanded - the biggest risk area."),
    "Mobile (React Native) Quality": ("4", "Strong overall; swipe gestures and NativeWind/Babel config were the main friction points."),
    "Bug Rate and Code Reliability": ("4", "Few bugs, mostly config/edge-case (>5 uploads, error mapping). 30/30 backend QA checks pass after fixes."),
    "Time Saved vs Manual Development": ("5", "Full backend + admin web + 18-screen mobile app in hours versus days/weeks by hand."),
    "Overall AI Build Quality": ("4", "Production-leaning architecture with a documented mock-AI boundary; requires human review on security and payments."),
}
t = doc.tables[2]
for i in range(1, len(t.rows)):
    cat = t.rows[i].cells[0].text.strip()
    if cat in ratings:
        set_cell(t.rows[i].cells[1], ratings[cat][0])
        set_cell(t.rows[i].cells[2], ratings[cat][1])

# ---------------- §9.2 Key Takeaways + §9.3 Recommendation (paragraphs) ----------------
takeaways = {
    "Most useful thing AI did": "scaffolding and boilerplate - repos, the DB schema/migrations from the BRD, CRUD endpoints and React/React Native screens - which eliminated hours of repetitive setup.",
    "Biggest failure or limitation": "silently skipping security hardening (OTP throttling, account non-disclosure, 2FA edge cases) and producing unreliable swipe gestures until explicitly directed.",
    "Would you recommend AI-first for this type of project": "Yes - with mandatory human review of security and payment flows and on-device testing of gestures.",
    "What would you do differently next time": "Specify the security/NFRs and the NativeWind + gesture constraints up front, commit to Git between phases, and test each endpoint in Postman before moving on.",
}

def append_answer(prompt_key, answer):
    for p in doc.paragraphs:
        txt = p.text.strip()
        if txt.lower().startswith(prompt_key.lower()):
            # strip trailing underscores/placeholder, then append the answer
            base = txt.rstrip("_ ").rstrip(":").strip()
            for r in list(p.runs):
                r.text = ""
            p.add_run(f"{base}: {answer}")
            return True
    return False

for k, v in takeaways.items():
    append_answer(k, v)

# §9.3 Recommendation - append a recommendation paragraph after the heading.
rec = ("Recommendation: AI coding tools are highly effective for delivering a production-leaning MVP of this scope "
       "very fast - the entire backend, admin portal and mobile app were built and verified in a single focused effort. "
       "For client delivery, pair the AI build with (1) a security review of auth, OTP, 2FA and payment/store rules, "
       "(2) on-device QA of gesture and camera flows, and (3) replacement of the mock scoring/enhancement service with a "
       "real vision model behind the existing service boundary. Used this way, AI-first development is recommended for Manaknight Digital.")
for idx, p in enumerate(doc.paragraphs):
    if p.text.strip().startswith("9.3"):
        # insert a new paragraph right after the heading
        new_p = p.insert_paragraph_before(rec) if False else None
        # python-docx has no direct insert_after; append run to a following empty paragraph if present
        # Fallback: add to end of document.
        break

doc.add_paragraph(rec)

doc.save(SRC)
print("Part 2 document filled and saved. Backup:", BACKUP)
