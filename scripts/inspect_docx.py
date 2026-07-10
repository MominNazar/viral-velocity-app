import sys
from docx import Document

path = sys.argv[1]
doc = Document(path)
print(f"Tables: {len(doc.tables)}")
for ti, t in enumerate(doc.tables):
    rows = len(t.rows)
    cols = len(t.columns)
    header = [c.text.strip() for c in t.rows[0].cells]
    print(f"\n--- Table {ti}: {rows} rows x {cols} cols ---")
    print("header:", header)
    # print first column of each row to identify the table
    firstcol = [t.rows[r].cells[0].text.strip() for r in range(min(rows, 25))]
    print("col0:", firstcol)
