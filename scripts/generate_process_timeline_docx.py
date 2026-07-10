#!/usr/bin/env python3
"""Generate Viral_Velocity_App_Process_Timeline.docx from project documentation."""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "Viral_Velocity_App_Process_Timeline.docx"


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            table.rows[ri + 1].cells[ci].text = str(val)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return table


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_code(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    p.paragraph_format.left_indent = Inches(0.25)


def build():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Title
    t = doc.add_heading("Viral Velocity App", 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph("Full Process Documentation & Timeline")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].bold = True
    sub.runs[0].font.size = Pt(14)
    meta = doc.add_paragraph(
        "Purpose: Document the AI-assisted build process for experience summary and management review.\n"
        "Audience: Manager / reviewer / Part 1 & Part 2 report authors.\n"
        "Last updated: June 2026"
    )
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # 1 Executive Summary
    doc.add_heading("1. Executive Summary", level=1)
    doc.add_paragraph(
        "The Viral Velocity Engine is a photo scoring and AI-enhancement product built as a monorepo "
        "with three applications:"
    )
    add_table(doc,
        ["App", "Stack", "Role"],
        [
            ("Backend", "Node.js, Express (ESM), SQLite", "REST API, auth, photos, scoring, enhancement, subscriptions, admin"),
            ("Admin Web", "React, Vite, TypeScript, React Query", "Internal portal for KPIs, users, discounts, moderation"),
            ("Mobile", "Expo SDK 56, React Native, NativeWind", "End-user app (upload, score, enhance, library, billing, settings)"),
        ],
        [1.2, 2.0, 3.3],
    )
    doc.add_paragraph("Build method: AI coding assistant (Cursor) with human direction, QA, and iteration.")
    doc.add_paragraph("Duration: ~5 planned days (core build) + post-delivery device testing and UI polish.")
    doc.add_paragraph(
        "Verification: Backend automated QA = 30/30 BRD-mapped checks; mobile tested on Android physical device via Expo Go on LAN."
    )

    # 2 Scope
    doc.add_heading("2. Project Scope (BRD Mapping)", level=1)
    doc.add_paragraph("The build follows Viral_Velocity_BRD.docx. High-level requirement groups:")
    add_table(doc,
        ["Area", "Key FRs", "Status"],
        [
            ("Auth & onboarding", "FR-1–6 (signup, login, OTP, parental gate, ToS, free trial)", "Complete"),
            ("Photo upload & scoring", "FR-7–11 (multi-upload, social score, sub-scores)", "Complete"),
            ("Enhancement", "FR-12–16 (5 versions, swipe, compare, prompts)", "Complete"),
            ("Library", "FR-18–20 (sort, filter, batch prompt)", "Complete"),
            ("Subscription", "FR-21–25 (plans, discounts, upgrade, cancel, paywall)", "Complete"),
            ("Settings", "FR-26–27 (profile, disable account)", "Complete"),
            ("Admin", "FR-28–36 (dashboard, images matched, subscribers, 2FA)", "Complete"),
        ],
        [1.8, 2.8, 1.0],
    )
    doc.add_paragraph(
        "BRD tension resolved: Subscription wireframe describes token tiers (Starter/Pro/Expert); "
        "FR-21–25 originally described Monthly/Annual. Final mobile UI implements wireframe tier + token model; "
        "backend supports tier pricing, token balance, one-time packs, and subscription lifecycle."
    )

    # 3 Timeline
    doc.add_heading("3. Timeline (Day-by-Day + Post-Delivery)", level=1)

    phases = [
        ("Phase 0 — Planning & Setup (Day 0)", [
            ("0.1", "Read Part 1 Overview + Part 2 Log templates + BRD", "Understood deliverables"),
            ("0.2", "Initialize monorepo (backend/, admin-web/, mobile/, docs/)", "Single repo, root README"),
            ("0.3", "Record architecture decisions", "docs/DECISIONS.md"),
        ]),
        ("Phase 1 — Backend Foundation (Day 1)", [
            ("AM", "Express scaffold, config, error middleware, SQLite", "backend/src/server.js, app.js, config.js"),
            ("AM", "Full DB migration from BRD MySQL schema", "backend/migrations/001_init.sql"),
            ("PM", "Auth: signup, login, Remember Me, OTP reset, rate limits", "backend/src/routes/auth.js"),
            ("PM", "Parental consent, welcome email, audit log", "auth.js, services/email.js"),
            ("PM", "JWT + bcrypt utilities", "backend/src/lib/tokens.js"),
        ]),
        ("Phase 2 — Backend Core + Admin APIs (Day 2)", [
            ("AM", "Photo upload (1–5), scoring, moderation gate", "routes/photos.js, services/ai.js"),
            ("AM", "Enhancement: 5 versions, swipe save/discard, prompts", "photos.js, services/ai.js"),
            ("PM", "Subscription: plans, subscribe, upgrade, cancel", "routes/subscriptions.js"),
            ("PM", "Admin: dashboard, images matched, subscribers, pricing, 2FA", "routes/admin.js"),
            ("PM", "Smoke script + QA harness", "backend/scripts/qa.mjs → 30 checks"),
        ]),
        ("Phase 3 — Admin Web Portal (Day 2–3)", [
            ("PM", "Vite + React + TS scaffold, API client, auth context", "admin-web/src/"),
            ("PM", "Login + 2FA challenge", "Login, Verify2FA"),
            ("PM", "Dashboard KPIs, Images Matched, Subscribers, Discount modal", "Dashboard, Subscribers"),
            ("PM", "Admin profile, change password, enable/disable 2FA", "Profile"),
        ]),
        ("Phase 4 — Mobile App (Day 3–4)", [
            ("AM", "Expo + NativeWind + React Query + navigation", "mobile/App.tsx"),
            ("AM", "Auth flow (6 screens)", "Landing through ResetPassword"),
            ("PM", "Core flow: Dashboard, PhotoUpload, Score", "Core screens"),
            ("PM", "Enhance swipe stack (gesture-handler + Reanimated)", "EnhanceScreen"),
            ("PM", "Compare, Library, ImageDetail, Subscription, Settings", "Remaining screens"),
        ]),
        ("Phase 5 — QA, Reports & Polish (Day 5)", [
            ("QA", "Run backend/scripts/qa.mjs", "30/30 PASS"),
            ("Docs", "Fill Part 2 Word doc", "Viral_Velocity_App_Part2_Logs.docx"),
            ("Docs", "Mirror logs in markdown", "docs/EXPERIENCE_LOG.md"),
            ("Mobile", "tsc --noEmit + Metro bundle", "Clean build"),
        ]),
    ]

    for title, rows in phases:
        doc.add_heading(title, level=2)
        add_table(doc, ["Time/Step", "Task", "Outcome"], rows, [0.8, 2.8, 2.9])

    doc.add_heading("Phase 6 — Real Device Testing & Iteration (Post Day 5)", level=2)
    doc.add_paragraph(
        "Live testing on Android phone + Expo Go + LAN backend (http://192.168.1.4:4000)."
    )
    add_table(doc,
        ["#", "Issue reported", "Root cause", "Fix"],
        [
            ("6.1", "Upload / keyboard / nav bar overlap", "Safe area insets = 0 on 3-button nav", "useScreenInsets() helper"),
            ("6.2", "OTP / login on device", "API base URL", "EXPO_PUBLIC_API_BASE for LAN IP"),
            ("6.3", "Add Prompt not working", "Replicate 402/429/422", "Default to Sharp (free local processing)"),
            ("6.4", "Enhance 500 error", "Sharp gamma() < 1.0", "Brightness-only darkening"),
            ("6.5", "Prompt 500 error", "result is not defined", "Fixed photos.js route"),
            ("6.6", "Prompts subtle / not working", "Keyword parser gaps", "Expanded analyzePrompt() + user notices"),
            ("6.7", "No paid APIs wanted", "Replicate requires credits", "ENHANCE_ENGINE=sharp in .env"),
            ("6.8", "Subscription wireframe", "Monthly-only UI", "Tier plans + token balance + Profile tab"),
            ("6.9", "Status bar / white nav bar", "Missing safe area + theme", "SafeAreaView top + dark app.json"),
            ("6.10", "App crash on launch", "Expo SDK 56 API removed", "NavigationBar.setStyle('dark')"),
            ("6.11", "Settings hard to find", "Hidden gear icon only", "Subscription | Settings toggle"),
        ],
        [0.4, 1.5, 1.8, 2.8],
    )
    doc.add_paragraph("Migration added: 002_tokens_and_tiers.sql (token_balance, plan_tier, billing_cycle).")

    # 4 Architecture
    doc.add_heading("4. System Architecture", level=1)
    arch = """MOBILE (Expo Go)
  Home | Gallery | Library | Profile (Subscription | Settings)
       ↓ HTTP /api (JWT Bearer)
BACKEND (Express :4000)
  /auth  /photos  /subscription  /admin
  SQLite + ai.js (mock scorer) + imageEnhance.js (Sharp / optional Replicate)
       ↓
ADMIN WEB (Vite :5173)
  Dashboard | Images Matched | Subscribers | Discount | 2FA"""
    add_code(doc, arch)

    doc.add_heading("Enhancement Engine", level=2)
    add_table(doc,
        ["Mode", "Config", "Cost", "Capability"],
        [
            ("Sharp (default)", "ENHANCE_ENGINE=sharp", "Free", "Color, contrast, blur, B&W — not object editing"),
            ("Replicate (optional)", "ENHANCE_ENGINE=replicate + token", "Paid credits", "Instruct-pix2pix style edits"),
        ],
        [1.2, 2.0, 0.8, 2.5],
    )

    # 5 Features
    doc.add_heading("5. Complete Feature Inventory", level=1)
    doc.add_heading("5.1 Backend API", level=2)
    add_bullets(doc, [
        "Auth: signup, login, forgot-password, OTP reset, profile, change-password, disable-profile",
        "Photos: upload, list, library, enhance (5 versions), prompt, swipe save/discard, delete",
        "Subscription: plans, subscribe, upgrade, one-time purchase, cancel",
        "Admin: login, 2FA, dashboard, images-matched, subscribers, pricing, retention, moderation",
    ])
    doc.add_heading("5.2 Mobile Screens (17)", level=2)
    add_table(doc,
        ["Screen group", "Purpose"],
        [
            ("Landing, Login, SignUp, ForgotPassword, Otp, ResetPassword", "Authentication"),
            ("Dashboard, Gallery", "Home and photo grid + token balance"),
            ("PhotoUpload, Score, Enhance, Compare, ImageDetail", "Upload, score, enhance workflow"),
            ("Library", "Saved images, batch prompt"),
            ("Subscription, Settings, ChangePassword", "Billing and account (Profile tab)"),
        ],
        [3.0, 3.5],
    )

    # 6 Decisions
    doc.add_heading("6. Key Architecture Decisions", level=1)
    add_table(doc,
        ["ID", "Decision", "Rationale"],
        [
            ("D-1", "SQLite over MySQL", "Zero external DB setup on Windows"),
            ("D-2", "Subscription lifecycle per FR-21–25", "Schema + QA consistency"),
            ("D-3", "Mock deterministic scorer", "No API key required"),
            ("D-4", "JWT + bcrypt, OTP rate limits", "NFR-6 security"),
            ("D-5", "Console email in dev", "Demo-friendly"),
            ("D-6", "Configurable moderation stub", "NFR-10 without vision API"),
            ("D-7", "Sharp as default enhancer", "Free, offline, no Replicate credits"),
            ("D-8", "Tier + token UI (wireframe)", "Matches SB-F QA + wireframe"),
            ("D-9", "Keyword prompt parser", "Responsive Sharp prompts"),
            ("D-10", "Subscription maintenance job", "Expire subs + purge data"),
        ],
        [0.5, 2.2, 3.8],
    )

    # 7 Bugs
    doc.add_heading("7. Bug Log Summary", level=1)
    add_table(doc,
        ["#", "Area", "Symptom", "Fix"],
        [
            ("1", "OTP", "No brute-force protection", "Rate limit middleware"),
            ("2", "Deps", "multer 1.x advisory", "Upgraded to 2.x"),
            ("3", "Forgot password", "Email enumeration", "Uniform API response"),
            ("4", "Upload", ">5 files → 500", "Map multer LIMIT → 400"),
            ("5", "Mobile swipe", "Janky gestures", "Reanimated + Pan gesture"),
            ("6", "Mobile build", "babel-preset-expo missing", "Added dependency"),
            ("7", "Enhance", "gamma() crash", "Brightness-only darkening"),
            ("8", "Prompt route", "result is not defined", "Fixed route vars"),
            ("9", "Replicate", "402/429/422", "Default to Sharp"),
            ("10", "Expo SDK 56", "nav bar API removed", "NavigationBar.setStyle"),
            ("11", "UI", "Status bar overlap", "Top safe area"),
            ("12", "UX", "Settings hidden", "Subscription | Settings toggle"),
        ],
        [0.4, 1.0, 2.0, 2.9],
    )

    # 8 Testing
    doc.add_heading("8. Testing & Verification", level=1)
    doc.add_heading("Automated (backend)", level=2)
    add_code(doc, "cd backend\nnpm run dev\nnode scripts/qa.mjs   # expect 30/30 PASS")
    doc.add_heading("Manual mobile checklist", level=2)
    add_bullets(doc, [
        "Signup / login on LAN",
        "Upload 1–5 photos → score screen",
        "Enhance → swipe save/discard",
        "Add Prompt (e.g. warm golden sunset, very vibrant)",
        "Profile → Subscription → select plan → Upgrade",
        "Profile → Settings → edit profile → Change Password",
        "Token balance visible on Home / Profile",
    ])

    # 9 AI Experience
    doc.add_heading("9. AI Assistance — Experience Summary", level=1)
    doc.add_heading("What AI did well (Score 4–5)", level=2)
    add_bullets(doc, [
        "Rapid scaffolding of Express + SQLite schema from BRD",
        "Consistent REST patterns, validation, JWT auth",
        "Admin web pages with React Query",
        "Mobile screen layout with NativeWind",
        "End-to-end smoke scripts and QA mapping",
    ])
    doc.add_heading("What needed human steering (Score 3–4)", level=2)
    add_bullets(doc, [
        "Security details (OTP rate limit, no email enumeration) — omitted unless explicitly requested",
        "React Native gestures / Reanimated — needed iteration",
        "BRD contradictions (token vs Monthly model) — needed product decision",
        "Real-device issues (LAN IP, safe areas, Android nav bar) — not caught in simulator",
        "Sharp vs Replicate trade-offs — needed cost/reality discussion",
    ])
    doc.add_heading("Overall assessment", level=2)
    p = doc.add_paragraph()
    p.add_run(
        "AI accelerated ~70–80% of boilerplate and CRUD. Highest value on backend and admin web. "
        "Mobile gestures, device-specific UI, and enhancement realism required the most manual QA and follow-up prompts. "
        "For production, AI is strong as a pair programmer but weak as a sole QA engineer on physical devices and paid API edge cases."
    ).italic = True

    # 10 How to run
    doc.add_heading("10. How to Run (Demo)", level=1)
    doc.add_heading("Start backend", level=2)
    add_code(doc, "cd backend\nnpm install\nnpm run migrate\nnpm run seed\nnpm run dev\n# LAN IP e.g. 192.168.1.4")
    doc.add_heading("Start mobile", level=2)
    add_code(doc, 'cd mobile\n$env:EXPO_PUBLIC_API_BASE="http://192.168.1.4:4000"\nnpx expo start --lan')
    doc.add_heading("Demo logins", level=2)
    add_table(doc,
        ["Role", "Email", "Password"],
        [
            ("Admin", "admin@viralvelocity.app", "Admin123"),
            ("User", "emily@example.com", "Demo1234"),
            ("QA test", "test@example.com", "Passw0rd"),
        ],
        [1.0, 2.5, 1.5],
    )

    # 11 Final summary
    doc.add_heading("11. Final State Summary", level=1)
    doc.add_paragraph(
        "We delivered a working Viral Velocity Engine monorepo: REST backend with SQLite, React admin portal, "
        "and Expo mobile app covering auth, photo upload/scoring, AI-style enhancement with user prompts, "
        "library management, tiered subscription/billing with token balance, and admin operations including "
        "discounts and 2FA. Development followed a 5-day AI-assisted plan, then extended with real-device testing "
        "that surfaced enhancement API costs, Sharp-based free processing, prompt parser improvements, subscription "
        "wireframe alignment, and Android UI fixes. Backend QA passes 30/30 automated checks; the mobile app runs "
        "on Expo Go over LAN. The process demonstrates that AI tools excel at scaffolding and CRUD but require "
        "human oversight for security, gestures, device layout, and third-party API economics."
    )

    doc.add_paragraph()
    end = doc.add_paragraph("— End of process documentation —")
    end.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(OUT)
    print(f"Created: {OUT}")


if __name__ == "__main__":
    build()
